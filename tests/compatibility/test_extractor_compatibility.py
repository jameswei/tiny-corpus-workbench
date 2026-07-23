from __future__ import annotations

import hashlib
import importlib.metadata
import json
import os
import socket
import tempfile
import unittest
from contextlib import contextmanager
from pathlib import Path
from unittest import mock

from docling.datamodel.base_models import InputFormat
from docling.datamodel.pipeline_options import (
    AcceleratorDevice,
    AcceleratorOptions,
    PdfPipelineOptions,
)
from docling.document_converter import DocumentConverter, PdfFormatOption
from docling_core.types.doc import (
    BoundingBox,
    DoclingDocument,
    DocItemLabel,
    ProvenanceItem,
    Size,
    TableCell,
    TableData,
)
from docx import Document
from markitdown import MarkItDown, StreamInfo
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen.canvas import Canvas

from tiny_corpus_workbench.runtime import RUNTIME_DEPENDENCIES


MODEL_ROOT_ENV = "TCW_DOCLING_ARTIFACTS"


@contextmanager
def network_denied():
    """Fail the spike if observation attempts to open a network connection."""

    def deny_connection(*args, **kwargs):
        raise AssertionError("observation-time network access attempted")

    with (
        mock.patch.object(socket, "create_connection", deny_connection),
        mock.patch.object(socket.socket, "connect", deny_connection),
        mock.patch.object(socket.socket, "connect_ex", deny_connection),
    ):
        yield


class ExtractorCompatibilitySpike(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        model_root = os.environ.get(MODEL_ROOT_ENV, ".cache/docling/models")
        cls.model_root = Path(model_root).resolve()
        if not cls.model_root.is_dir():
            raise unittest.SkipTest(
                f"prefetch Docling models and set {MODEL_ROOT_ENV}: {cls.model_root}"
            )

        cls.temp = tempfile.TemporaryDirectory()
        cls.root = Path(cls.temp.name)
        cls.inputs = cls._write_inputs(cls.root)

    @classmethod
    def tearDownClass(cls) -> None:
        if hasattr(cls, "temp"):
            cls.temp.cleanup()

    @staticmethod
    def _write_inputs(root: Path) -> dict[str, Path]:
        text = "Compatibility Spike\n\nTCW-SPIKE-01\n"

        pdf = root / "spike.pdf"
        canvas = Canvas(str(pdf), pagesize=letter, invariant=1)
        canvas.setTitle("Compatibility Spike")
        canvas.drawString(72, 720, "Compatibility Spike")
        canvas.drawString(72, 700, "TCW-SPIKE-01")
        canvas.save()

        docx = root / "spike.docx"
        document = Document()
        document.add_heading("Compatibility Spike", level=1)
        document.add_paragraph("TCW-SPIKE-01")
        document.save(docx)

        markdown = root / "spike.md"
        markdown.write_text("# Compatibility Spike\n\nTCW-SPIKE-01\n", "utf-8")

        plain = root / "spike.txt"
        plain.write_text(text, "utf-8")
        return {"pdf": pdf, "docx": docx, "markdown": markdown, "text": plain}

    def _docling_converter(self) -> DocumentConverter:
        options = PdfPipelineOptions(
            artifacts_path=self.model_root,
            do_ocr=False,
            do_table_structure=True,
            enable_remote_services=False,
            allow_external_plugins=False,
            accelerator_options=AcceleratorOptions(device=AcceleratorDevice.CPU),
        )
        return DocumentConverter(
            allowed_formats=list(InputFormat),
            format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=options)},
        )

    def test_exact_versions_and_all_extraction_views_offline(self) -> None:
        expected_versions = {
            **RUNTIME_DEPENDENCIES,
            "reportlab": "5.0.0",
            "python-docx": "1.2.0",
            "jsonschema": "4.26.0",
        }
        self.assertEqual(
            {name: importlib.metadata.version(name) for name in expected_versions},
            expected_versions,
        )

        converter = self._docling_converter()
        markitdown = MarkItDown(enable_plugins=False)
        stream_info = {
            "markdown": StreamInfo(extension=".md", mimetype="text/markdown", charset="utf-8"),
            "text": StreamInfo(extension=".txt", mimetype="text/plain", charset="utf-8"),
        }
        schema_identities: set[tuple[str, str]] = set()

        with network_denied():
            for name, source in self.inputs.items():
                with self.subTest(format=name):
                    result = converter.convert(source)
                    doc = result.document
                    json_path = self.root / f"{name}.json"
                    markdown_path = self.root / f"{name}.docling.md"
                    doc.save_as_json(json_path)
                    doc.save_as_markdown(markdown_path)

                    payload = json.loads(json_path.read_text("utf-8"))
                    self.assertGreater(json_path.stat().st_size, 0)
                    self.assertTrue(markdown_path.read_text("utf-8").strip())
                    reloaded = DoclingDocument.load_from_json(json_path)
                    self.assertEqual(reloaded.export_to_markdown(), doc.export_to_markdown())
                    traversed = list(reloaded.iterate_items(with_groups=True))
                    self.assertTrue(traversed)
                    for item, _ in traversed:
                        for child in item.children:
                            self.assertIsNotNone(child.resolve(reloaded))
                        for provenance in getattr(item, "prov", []):
                            self.assertGreater(provenance.page_no, 0)
                            self.assertIsNotNone(provenance.bbox)
                    for page in reloaded.pages.values():
                        self.assertGreater(page.size.width, 0)
                        self.assertGreater(page.size.height, 0)
                    schema_identities.add(
                        (str(payload.get("schema_name")), str(payload.get("version")))
                    )

                    kwargs = {}
                    if name in stream_info:
                        kwargs["stream_info"] = stream_info[name]
                    converted = markitdown.convert_local(source, **kwargs)
                    self.assertTrue(converted.markdown.strip())

        self.assertEqual(schema_identities, {("DoclingDocument", "1.10.0")})
        inventory = []
        for path in sorted(self.model_root.rglob("*")):
            if path.is_symlink():
                self.fail(f"model inventory contains symlink: {path}")
            if path.is_file():
                inventory.append(
                    {
                        "path": path.relative_to(self.model_root).as_posix(),
                        "size": path.stat().st_size,
                        "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
                    }
                )
        self.assertTrue(inventory, "model inventory must contain regular files")

    def test_canonical_relationship_table_and_geometry_apis(self) -> None:
        document = DoclingDocument(name="diagnosis-compatibility")
        document.add_page(1, Size(width=612, height=792))
        provenance = ProvenanceItem(
            page_no=1,
            bbox=BoundingBox(l=72, t=72, r=180, b=90),
            charspan=(0, 7),
        )
        caption = document.add_text(
            DocItemLabel.CAPTION, "Caption", prov=provenance
        )
        table = document.add_table(
            TableData(
                table_cells=[
                    TableCell(
                        start_row_offset_idx=0,
                        end_row_offset_idx=1,
                        start_col_offset_idx=0,
                        end_col_offset_idx=1,
                        text="Cell",
                    )
                ],
                num_rows=1,
                num_cols=1,
            ),
            caption=caption,
            prov=provenance,
        )
        path = self.root / "relationship.json"
        document.save_as_json(path)
        reloaded = DoclingDocument.load_from_json(path)
        self.assertEqual(reloaded.tables[0].data.table_cells[0].text, "Cell")
        self.assertEqual(
            reloaded.tables[0].captions[0].resolve(reloaded).self_ref,
            caption.self_ref,
        )
        self.assertEqual(table.self_ref, "#/tables/0")
        self.assertEqual(reloaded.pages[1].size, Size(width=612, height=792))


if __name__ == "__main__":
    unittest.main()
