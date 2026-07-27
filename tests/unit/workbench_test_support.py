from __future__ import annotations

import json
import io
import os
import subprocess
import sys
import tempfile
from contextlib import redirect_stderr, redirect_stdout
from pathlib import Path
from unittest import mock

from tiny_corpus_workbench import cli
from tiny_corpus_workbench.artifacts import canonical_json


REPOSITORY = Path(__file__).resolve().parents[2]


def run_tcw(*arguments: str) -> dict:
    completed = subprocess.run(
        [sys.executable, "-m", "tiny_corpus_workbench", *arguments],
        cwd=REPOSITORY,
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(completed.stdout)


class PublishedObservation:
    def __init__(self) -> None:
        self.temporary = tempfile.TemporaryDirectory(
            dir=Path(tempfile.gettempdir()).resolve()
        )
        result = run_tcw(
            "observe",
            str(REPOSITORY / "fixtures/golden/policy-memo.md"),
            "--output-root",
            self.temporary.name,
        )
        self.root = Path(result["manifest"]).parent

    def close(self) -> None:
        self.temporary.cleanup()


class PublishedDiagnosis(PublishedObservation):
    def __init__(self) -> None:
        super().__init__()
        result = run_tcw(
            "diagnose",
            str(self.root),
            "--output-root",
            str(Path(self.temporary.name) / "diagnoses"),
        )
        self.diagnosis = Path(result["manifest"]).parent


class PublishedFailedObservation:
    def __init__(self) -> None:
        self.temporary = tempfile.TemporaryDirectory(
            dir=Path(tempfile.gettempdir()).resolve()
        )
        stdout, stderr = io.StringIO(), io.StringIO()
        with (
            mock.patch(
                "tiny_corpus_workbench.extractors.docling.convert",
                side_effect=RuntimeError("stable docling failure"),
            ),
            mock.patch(
                "tiny_corpus_workbench.extractors.markitdown.convert",
                side_effect=RuntimeError("stable markitdown failure"),
            ),
            redirect_stdout(stdout),
            redirect_stderr(stderr),
        ):
            code = cli.main(
                [
                    "observe",
                    str(REPOSITORY / "fixtures/golden/policy-memo.md"),
                    "--output-root",
                    self.temporary.name,
                ]
            )
        if code != 4:
            raise RuntimeError(stderr.getvalue())
        self.root = Path(json.loads(stdout.getvalue())["manifest"]).parent

    def close(self) -> None:
        self.temporary.cleanup()


class PublishedCorpus:
    def __init__(self) -> None:
        self.temporary = tempfile.TemporaryDirectory(
            dir=Path(tempfile.gettempdir()).resolve()
        )
        temporary_root = Path(self.temporary.name)
        input_root = temporary_root / "input"
        input_root.mkdir()
        sources = (
            ("meeting-minutes", "fixtures/golden/meeting-minutes.md"),
            ("policy-memo", "fixtures/golden/policy-memo.md"),
            ("release-notice", "fixtures/golden/release-notice.md"),
            ("short-note", "fixtures/diagnosis/short-note.md"),
            (
                "structural-traps",
                "fixtures/diagnosis/structural-traps.md",
            ),
        )
        corpus_spec = input_root / "model-free-corpus.json"
        corpus_spec.write_bytes(
            canonical_json(
                {
                    "schema_version": "tcw.corpus-spec/v0.5",
                    "corpus_id": "model-free-workbench-corpus",
                    "title": "Model-free workbench corpus",
                    "members": [
                        {
                            "member_id": member_id,
                            "family": member_id,
                            "format": "md",
                            "source": os.path.relpath(
                                REPOSITORY / relative_source,
                                corpus_spec.parent,
                            ),
                        }
                        for member_id, relative_source in sources
                    ],
                }
            )
        )
        result = run_tcw(
            "inspect-corpus",
            str(corpus_spec),
            "--output-root",
            str(temporary_root / "output"),
            "--docling-artifacts",
            str(temporary_root / "missing-models"),
        )
        self.root = Path(result["manifest"]).parent

    def close(self) -> None:
        self.temporary.cleanup()


class PublishedRefinements:
    def __init__(self) -> None:
        self.temporary = tempfile.TemporaryDirectory(
            dir=Path(tempfile.gettempdir()).resolve()
        )
        result = run_tcw(
            "observe",
            str(REPOSITORY / "fixtures/refinement/v0.5/whitespace-cleanup.md"),
            "--output-root",
            self.temporary.name,
        )
        self.observation = Path(result["manifest"]).parent
        result = run_tcw(
            "diagnose",
            str(self.observation),
            "--output-root",
            str(Path(self.temporary.name) / "diagnoses"),
        )
        self.diagnosis = Path(result["manifest"]).parent
        findings = json.loads((self.diagnosis / "findings.json").read_text("utf-8"))
        finding_id = next(
            item["finding_id"]
            for item in findings["findings"]
            if item["rule_id"] == "TCW-D009"
        )
        self.applied = self._resolve(finding_id, "APPROVED", "applied")
        self.rejected = self._resolve(finding_id, "REJECTED", "rejected")

    def _resolve(self, finding_id: str, state: str, label: str) -> Path:
        decision_path = Path(self.temporary.name) / f"{label}-decision.json"
        run_tcw(
            "draft-refinement",
            str(self.diagnosis),
            "--finding",
            finding_id,
            "--base",
            str(self.observation),
            "--output",
            str(decision_path),
        )
        decision = json.loads(decision_path.read_text("utf-8"))
        decision["decision"] = {
            "state": state,
            "decided_by": "test-owner",
            "note": None,
        }
        decision_path.write_bytes(canonical_json(decision))
        result = run_tcw(
            "resolve-refinement",
            str(decision_path),
            "--diagnosis",
            str(self.diagnosis),
            "--base",
            str(self.observation),
            "--output-root",
            str(Path(self.temporary.name) / label),
        )
        return Path(result["manifest"]).parent

    def close(self) -> None:
        self.temporary.cleanup()


class PublishedChain:
    def __init__(self) -> None:
        from docx import Document

        self.temporary = tempfile.TemporaryDirectory(
            dir=Path(tempfile.gettempdir()).resolve()
        )
        source = Path(self.temporary.name) / "chain.docx"
        document = Document()
        document.add_heading("Chain fixture", level=1)
        paragraph = document.add_paragraph()
        paragraph.add_run(
            "This paragraph contains a non-breaking\u00a0space and   repeated "
            "ASCII spaces. It also contains enough stable project-authored "
            "text to avoid the short-document diagnosis while preserving a "
            "second independent finding for the next revision. The final inter-"
        )
        paragraph.add_run().add_break()
        paragraph.add_run(
            "national example remains available for deterministic repair."
        )
        document.save(source)
        observed = run_tcw(
            "observe",
            str(source),
            "--output-root",
            str(Path(self.temporary.name) / "observations"),
        )
        self.observation = Path(observed["manifest"]).parent
        first_diagnosis = self._diagnose(self.observation, "diagnosis-1")
        first_finding = self._finding(first_diagnosis, "TCW-D009")
        self.first = self._resolve(
            first_diagnosis,
            self.observation,
            first_finding,
            "revision-1",
        )
        second_diagnosis = self._diagnose(self.first, "diagnosis-2")
        second_finding = self._finding(second_diagnosis, "TCW-D010")
        self.second = self._resolve(
            second_diagnosis,
            self.first,
            second_finding,
            "revision-2",
        )
        self.first_diagnosis = first_diagnosis
        self.second_diagnosis = second_diagnosis

    def _diagnose(self, base: Path, label: str) -> Path:
        result = run_tcw(
            "diagnose",
            str(base),
            "--output-root",
            str(Path(self.temporary.name) / label),
        )
        return Path(result["manifest"]).parent

    @staticmethod
    def _finding(diagnosis: Path, rule_id: str) -> str:
        findings = json.loads((diagnosis / "findings.json").read_text("utf-8"))
        return next(
            item["finding_id"]
            for item in findings["findings"]
            if item["rule_id"] == rule_id
        )

    def _resolve(
        self, diagnosis: Path, base: Path, finding_id: str, label: str
    ) -> Path:
        decision_path = Path(self.temporary.name) / f"{label}.json"
        run_tcw(
            "draft-refinement",
            str(diagnosis),
            "--finding",
            finding_id,
            "--base",
            str(base),
            "--output",
            str(decision_path),
        )
        decision = json.loads(decision_path.read_text("utf-8"))
        decision["decision"] = {
            "state": "APPROVED",
            "decided_by": "chain-owner",
            "note": None,
        }
        decision_path.write_bytes(canonical_json(decision))
        result = run_tcw(
            "resolve-refinement",
            str(decision_path),
            "--diagnosis",
            str(diagnosis),
            "--base",
            str(base),
            "--output-root",
            str(Path(self.temporary.name) / f"{label}-output"),
        )
        return Path(result["manifest"]).parent

    def close(self) -> None:
        self.temporary.cleanup()
