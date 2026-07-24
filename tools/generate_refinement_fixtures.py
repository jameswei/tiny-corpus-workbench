#!/usr/bin/env python3
"""Check the deterministic project-authored v0.3 refinement fixtures."""

from __future__ import annotations

import argparse
import hashlib
import json
import tempfile
import zipfile
from datetime import UTC, datetime
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
FIXTURES = ROOT / "fixtures/refinement/v0.3"
REGISTRY = FIXTURES / "fixtures.json"
HYPHENATION = FIXTURES / "line-end-hyphenation.docx"
WHITESPACE = FIXTURES / "whitespace-cleanup.md"
FIXED_ZIP_TIME = (2026, 7, 24, 0, 0, 0)
FIXED_DOC_TIME = datetime(2026, 7, 24, tzinfo=UTC)
WHITESPACE_BYTES = (
    "# Whitespace Cleanup Fixture\n\n"
    "This line contains a non-breaking\u00a0space and   repeated ASCII spaces. "
    "The remaining project-authored text is deliberately long enough to keep "
    "the suspiciously-short rule out of this fixture. It preserves stable "
    "evidence for deterministic whitespace diagnosis and controlled revision "
    "practice.\n"
).encode()


def write_hyphenation(path: Path) -> None:
    document = Document()
    document.core_properties.title = "Line-End Hyphenation Fixture"
    document.core_properties.author = "tiny-corpus-workbench contributors"
    document.core_properties.created = FIXED_DOC_TIME
    document.core_properties.modified = FIXED_DOC_TIME
    paragraph = document.add_paragraph()
    paragraph.add_run("Inter-")
    paragraph.add_run().add_break()
    paragraph.add_run(
        "operable systems preserve deterministic evidence across prepared "
        "revisions. This project-authored fixture is deliberately long enough "
        "to avoid the suspiciously-short diagnosis. Additional stable content "
        "keeps the single intended line-end hyphenation candidate inspectable."
    )
    temporary = path.with_suffix(".raw.docx")
    document.save(temporary)
    with zipfile.ZipFile(temporary) as source, zipfile.ZipFile(
        path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
    ) as target:
        for name in sorted(source.namelist()):
            info = zipfile.ZipInfo(name, FIXED_ZIP_TIME)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            info.create_system = 3
            target.writestr(
                info,
                source.read(name),
                compress_type=zipfile.ZIP_DEFLATED,
                compresslevel=9,
            )
    temporary.unlink()


def descriptor(path: Path, fixture_id: str, rules: list[str]) -> dict[str, object]:
    raw = path.read_bytes()
    return {
        "id": fixture_id,
        "path": path.relative_to(ROOT).as_posix(),
        "size": len(raw),
        "sha256": hashlib.sha256(raw).hexdigest(),
        "license": "CC0-1.0",
        "expected_rules": rules,
    }


def expected() -> dict[str, object]:
    return {
        "schema_version": "tcw.refinement-fixture-registry/v0.3",
        "generator": "tools/generate_refinement_fixtures.py",
        "fixtures": [
            descriptor(
                HYPHENATION,
                "line-end-hyphenation",
                ["TCW-D010"],
            ),
            descriptor(
                WHITESPACE,
                "whitespace-cleanup",
                ["TCW-D009"],
            ),
        ],
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--check", action="store_true")
    arguments = parser.parse_args()
    if arguments.check:
        with tempfile.TemporaryDirectory() as directory:
            generated = Path(directory) / HYPHENATION.name
            write_hyphenation(generated)
            if (
                not HYPHENATION.is_file()
                or HYPHENATION.read_bytes() != generated.read_bytes()
                or WHITESPACE.read_bytes() != WHITESPACE_BYTES
            ):
                print("v0.3 refinement fixture bytes differ")
                return 1
    else:
        write_hyphenation(HYPHENATION)
        WHITESPACE.write_bytes(WHITESPACE_BYTES)
    rendered = (
        json.dumps(expected(), ensure_ascii=False, indent=2, sort_keys=True) + "\n"
    ).encode()
    if arguments.check:
        if not REGISTRY.is_file() or REGISTRY.read_bytes() != rendered:
            print("v0.3 refinement fixture registry differs")
            return 1
        return 0
    REGISTRY.write_bytes(rendered)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
