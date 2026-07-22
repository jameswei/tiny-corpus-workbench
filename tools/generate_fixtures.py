from __future__ import annotations

import argparse
import hashlib
import json
import tempfile
import zipfile
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen.canvas import Canvas


ROOT = Path(__file__).resolve().parents[1]
AUTHORED = ROOT / "fixtures/authored"
GOLDEN = ROOT / "fixtures/golden"
FORMATS = ("docx", "md", "pdf", "txt")
MEDIA_TYPES = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "md": "text/markdown",
    "txt": "text/plain",
}
FIXED_ZIP_TIME = (2026, 7, 20, 0, 0, 0)
FIXED_DOC_TIME = datetime(2026, 7, 20, tzinfo=UTC)


def digest(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def markdown(spec: dict[str, Any]) -> str:
    lines = [f"# {spec['title']}", "", f"Document ID: {spec['document_id']}", f"Date: {spec['date']}", f"Reference: {spec['url']}"]
    for section in spec["sections"]:
        lines.extend(["", f"## {section['heading']}"])
        for value in section.get("paragraphs", []):
            lines.extend(["", value])
        if section.get("bullets"):
            lines.append("")
            lines.extend(f"- {value}" for value in section["bullets"])
        if section.get("numbered"):
            lines.append("")
            lines.extend(f"{index}. {value}" for index, value in enumerate(section["numbered"], 1))
        if section.get("table"):
            table = section["table"]
            lines.extend(["", "| " + " | ".join(table[0]) + " |", "| " + " | ".join("---" for _ in table[0]) + " |"])
            lines.extend("| " + " | ".join(row) + " |" for row in table[1:])
    return "\n".join(lines) + "\n"


def plain_text(spec: dict[str, Any]) -> str:
    lines = [spec["title"].upper(), f"Document ID: {spec['document_id']}", f"Date: {spec['date']}", f"Reference: {spec['url']}"]
    for section in spec["sections"]:
        lines.extend(["", section["heading"].upper()])
        lines.extend(section.get("paragraphs", []))
        lines.extend(f"- {value}" for value in section.get("bullets", []))
        lines.extend(f"{index}. {value}" for index, value in enumerate(section.get("numbered", []), 1))
        if section.get("table"):
            widths = [max(len(row[index]) for row in section["table"]) for index in range(3)]
            lines.extend("  ".join(value.ljust(widths[index]) for index, value in enumerate(row)).rstrip() for row in section["table"])
    return "\n".join(lines) + "\n"


def _hyperlink(paragraph, url: str) -> None:
    relationship = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    properties.append(color)
    run.append(properties)
    text = OxmlElement("w:t")
    text.text = url
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def write_docx(spec: dict[str, Any], path: Path) -> None:
    doc = Document()
    doc.core_properties.title = spec["title"]
    doc.core_properties.author = "tiny-corpus-workbench contributors"
    doc.core_properties.created = FIXED_DOC_TIME
    doc.core_properties.modified = FIXED_DOC_TIME
    doc.add_heading(spec["title"], 0)
    doc.add_paragraph(f"Document ID: {spec['document_id']}")
    doc.add_paragraph(f"Date: {spec['date']}")
    reference = doc.add_paragraph("Reference: ")
    _hyperlink(reference, spec["url"])
    for section in spec["sections"]:
        doc.add_heading(section["heading"], 1)
        for value in section.get("paragraphs", []):
            doc.add_paragraph(value)
        for value in section.get("bullets", []):
            doc.add_paragraph(value, style="List Bullet")
        for value in section.get("numbered", []):
            doc.add_paragraph(value, style="List Number")
        if section.get("table"):
            table = doc.add_table(rows=0, cols=3)
            table.style = "Table Grid"
            for row in section["table"]:
                cells = table.add_row().cells
                for index, value in enumerate(row):
                    cells[index].text = value
    temporary = path.with_suffix(".raw.docx")
    doc.save(temporary)
    with zipfile.ZipFile(temporary) as source, zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as target:
        for name in sorted(source.namelist()):
            info = zipfile.ZipInfo(name, FIXED_ZIP_TIME)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            info.create_system = 3
            target.writestr(info, source.read(name), compress_type=zipfile.ZIP_DEFLATED, compresslevel=9)
    temporary.unlink()


def write_pdf(spec: dict[str, Any], path: Path) -> None:
    canvas = Canvas(str(path), pagesize=letter, invariant=1, pageCompression=1)
    canvas.setAuthor("tiny-corpus-workbench contributors")
    canvas.setCreator("tools/generate_fixtures.py")
    canvas.setTitle(spec["title"])
    y = 750

    def line(value: str, font: str = "Helvetica", size: int = 10, indent: int = 0) -> None:
        nonlocal y
        if y < 54:
            canvas.showPage()
            y = 750
        canvas.setFont(font, size)
        canvas.drawString(54 + indent, y, value)
        y -= size + 5

    line(spec["title"], "Helvetica-Bold", 16)
    line(f"Document ID: {spec['document_id']}")
    line(f"Date: {spec['date']}")
    line(f"Reference: {spec['url']}")
    canvas.linkURL(spec["url"], (54, y + 10, 400, y + 24), relative=0)
    for section in spec["sections"]:
        y -= 4
        line(section["heading"], "Helvetica-Bold", 12)
        for value in section.get("paragraphs", []):
            line(value)
        for value in section.get("bullets", []):
            line(f"- {value}", indent=12)
        for index, value in enumerate(section.get("numbered", []), 1):
            line(f"{index}. {value}", indent=12)
        if section.get("table"):
            table = section["table"]
            row_height = 22
            column_widths = (170, 120, 100)
            left = 66
            top = y
            bottom = top - row_height * len(table)
            canvas.setLineWidth(1)
            for index in range(len(table) + 1):
                row_y = top - index * row_height
                canvas.line(left, row_y, left + sum(column_widths), row_y)
            x = left
            canvas.line(x, top, x, bottom)
            for width in column_widths:
                x += width
                canvas.line(x, top, x, bottom)
            canvas.setFont("Helvetica", 9)
            for row_index, row in enumerate(table):
                x = left
                for column_index, value in enumerate(row):
                    if row_index == 0:
                        canvas.setFont("Helvetica-Bold", 9)
                    else:
                        canvas.setFont("Helvetica", 9)
                    canvas.drawString(x + 4, top - (row_index + 1) * row_height + 7, value)
                    x += column_widths[column_index]
            y = bottom - 8
    canvas.save()


def generate_into(output: Path) -> dict[str, Any]:
    output.mkdir(parents=True, exist_ok=True)
    lock_hash = digest(ROOT / "uv.lock")
    fixtures = []
    for authored_path in sorted(AUTHORED.glob("*.json")):
        spec = json.loads(authored_path.read_text("utf-8"))
        authored_hash = digest(authored_path)
        family = spec["family"]
        for format_name in FORMATS:
            path = output / f"{family}.{format_name}"
            if format_name == "md":
                path.write_text(markdown(spec), "utf-8", newline="\n")
            elif format_name == "txt":
                path.write_text(plain_text(spec), "utf-8", newline="\n")
            elif format_name == "docx":
                write_docx(spec, path)
            else:
                write_pdf(spec, path)
            fixtures.append({
                "id": f"{family}-{format_name}",
                "family": family,
                "format": format_name,
                "path": f"fixtures/golden/{path.name}",
                "media_type": MEDIA_TYPES[format_name],
                "size": path.stat().st_size,
                "sha256": digest(path),
                "authored_source": {"id": family, "path": f"fixtures/authored/{authored_path.name}", "sha256": authored_hash},
                "generator": {"name": "tools/generate_fixtures.py", "schema": spec["schema_version"], "lockfile_sha256": lock_hash},
                "ownership": "project-authored",
                "license": "CC0-1.0",
                "anchors": {"document_id": spec["document_id"], "date": spec["date"], "url": spec["url"]},
                "expected_docling_table_count": 0 if format_name == "txt" else 1,
            })
    fixtures.sort(key=lambda item: item["id"])
    registry = {"schema_version": "tcw.fixture-registry/v0.1", "generator": {"name": "tools/generate_fixtures.py", "schema": "tcw.authored-fixture/v0.1"}, "lockfile_sha256": lock_hash, "fixtures": fixtures}
    (output / "fixtures.json").write_text(json.dumps(registry, ensure_ascii=False, sort_keys=True, indent=2) + "\n", "utf-8", newline="\n")
    return registry


def snapshot(root: Path) -> dict[str, bytes]:
    return {path.name: path.read_bytes() for path in sorted(root.iterdir()) if path.is_file()}


def main() -> int:
    arguments = argparse.ArgumentParser()
    arguments.add_argument("--check", action="store_true")
    args = arguments.parse_args()
    if not args.check:
        generate_into(GOLDEN)
        return 0
    with tempfile.TemporaryDirectory() as first, tempfile.TemporaryDirectory() as second:
        generate_into(Path(first))
        generate_into(Path(second))
        generated = snapshot(Path(first))
        if generated != snapshot(Path(second)):
            raise SystemExit("fixture generation is not byte-deterministic")
        if generated != snapshot(GOLDEN):
            raise SystemExit("committed golden fixtures do not match deterministic generation")
    print("fixture generation is byte-deterministic and matches committed files")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
